# healthcare_doc_intelligence.py
import streamlit as st
import tempfile
import os
from pathlib import Path
from typing import List
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from langchain_community.llms import Ollama
from langchain_community.retrievers import PubMedRetriever
from sentence_transformers import SentenceTransformer
from langchain.embeddings.base import Embeddings
from langchain_community.vectorstores import Chroma
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

st.set_page_config(page_title="Healthcare Document Intelligence", layout="wide")

# --- UI Styling ---
def render_styles():
    st.markdown("""
    <style>
        .main-header { background: linear-gradient(135deg, #2c3e50, #3498db); color: white; padding: 2rem 0; border-radius: 8px; text-align: center; }
        .response-container { border-left: 4px solid #3498db; background: #f8f9fa; padding: 1.5rem; border-radius: 8px; margin: 1rem 0; }
        .stButton > button { background-color: #2c3e50; color: white; border-radius: 8px; }
    </style>
    """, unsafe_allow_html=True)

render_styles()

# --- Embedding (BioBERT / ClinicalBERT preferred if downloaded; else fallback) ---
class MedicalEmbeddings(Embeddings):
    def __init__(self, model_name="all-MiniLM-L6-v2"):
        try:
            self.model = SentenceTransformer("emilyalsentzer/Bio_ClinicalBERT")
        except Exception:
            self.model = SentenceTransformer(model_name)
    def embed_documents(self, texts: List[str]):
        return self.model.encode(texts, convert_to_tensor=False)
    def embed_query(self, text: str):
        return self.model.encode([text], convert_to_tensor=False)[0]

medical_embeddings = MedicalEmbeddings()

# --- PDF & DOCX Loader ---
def process_pdf(file_path):
    reader = PdfReader(file_path)
    all_text = ''
    for page in reader.pages:
        txt = page.extract_text() if page.extract_text() else ""
        all_text += txt + "\n"
    return all_text

def process_docx(file_path):
    doc = DocxDocument(file_path)
    all_text = "\n".join([para.text for para in doc.paragraphs])
    return all_text

def chunk_documents(text, source_name):
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
    return [Document(page_content=c, metadata={'source': source_name}) for c in splitter.split_text(text)]

# --- Vector Database ---
data_dir = Path("data/vectorstore")
data_dir.mkdir(parents=True, exist_ok=True)
vectorstore = Chroma(
    persist_directory=str(data_dir),
    embedding_function=medical_embeddings,
    collection_name="medical_documents"
)

# --- LLM (Ollama CLI must be running and model downloaded) ---
llm = Ollama(model="llama3.1:8b", base_url="http://localhost:11434", temperature=0.1)

# --- PubMed Retriever ---
pubmed_retriever = PubMedRetriever(doc_content_chars_max=2000, max_retry=3)

# --- Main UI ---
st.markdown('<div class="main-header"><h1>Healthcare Document Intelligence System</h1><p>AI-Powered Retrieval for Medical Professionals</p></div>', unsafe_allow_html=True)

with st.sidebar:
    st.header("Upload Medical Documents")
    uploaded_files = st.file_uploader("PDF / DOCX / TXT Only", type=["pdf", "docx", "txt"], accept_multiple_files=True)
    if st.button("Process & Index Documents") and uploaded_files:
        for file in uploaded_files:
            with tempfile.NamedTemporaryFile(delete=False) as tmp:
                tmp.write(file.read())
                tmp_path = tmp.name
                if file.name.endswith(".pdf"):
                    text = process_pdf(tmp_path)
                elif file.name.endswith(".docx"):
                    text = process_docx(tmp_path)
                else:
                    text = Path(tmp_path).read_text()
                docs = chunk_documents(text, file.name)
                vectorstore.add_documents(docs)
        st.success(f"{len(uploaded_files)} documents processed & indexed!")

# --- Query Section ---
st.subheader("Medical Knowledge Query")
medical_query = st.text_area("Enter your query (medical, clinical, drug, research)...", height=80)
mode = st.radio("Search Mode", ["Comprehensive", "Local Documents Only", "PubMed Only"])
max_results = st.slider("Results to show", 3, 10, 5)

if st.button("Search Medical Knowledge"):
    # Retrieve from local indexed docs
    local_docs = []
    if mode != "PubMed Only":
        local_docs = vectorstore.similarity_search(medical_query, k=max_results)

    # Retrieve from PubMed (real-time)
    pubmed_docs = []
    if mode in ["Comprehensive", "PubMed Only"]:
        pubmed_docs = pubmed_retriever.get_relevant_documents(medical_query)[:max_results]

    # --- Compose Context for LLM ---
    context = ""
    if local_docs:
        context += "\n\n".join([f"[{d.metadata.get('source')}] {d.page_content[:600]}" for d in local_docs])
    if pubmed_docs:
        context += "\n\n".join([f"[PubMed] {d.page_content[:600]}" for d in pubmed_docs])

    prompt = f"""
You are a medical AI assistant. Refer to the following context from clinical documents and real-time medical literature. 
Answers must be evidence-based, cite the source, and include medical disclaimer.
---
Context:
{context}

Query:
{medical_query}

Instructions:
1. Provide accurate, clinical-grade answers using only the context.
2. For facts use [SOURCE_NAME] to cite where info came from.
3. If insufficient info, clearly state that.
4. Include a medical disclaimer at the end.

Answer:
"""

    # --- LLM Inference ---
    st.markdown('<div class="response-container"><b>Response:</b></div>', unsafe_allow_html=True)
    try:
        response = llm.invoke(prompt)
        st.markdown(f"<div class='response-container'>{response}</div>", unsafe_allow_html=True)
    except Exception as e:
        st.error(f"LLM error: {e}")

    # --- Sources Display ---
    if local_docs:
        st.markdown("**Local Sources Cited:**")
        for d in local_docs:
            st.markdown(f"[{d.metadata.get('source')}] {d.page_content[:180]}...", unsafe_allow_html=False)

    if pubmed_docs:
        st.markdown("**PubMed Sources Cited:**")
        for i, d in enumerate(pubmed_docs, 1):
            st.markdown(f"**{i}. PubMed Abstract:** {d.page_content[:220]}...", unsafe_allow_html=False)

    # --- Disclaimer ---
    st.warning("This answer is for educational purposes only. Consult healthcare professionals for any medical decisions.")

# --- End ---
st.caption("© 2025 Healthcare Document Intelligence System — Built for Medical UX, Privacy, and Real-Time Evidence.")
